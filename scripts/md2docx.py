"""
Lightweight markdown -> .docx converter for the Entrium TZ.

Why custom (not pandoc): Windows machine has no pandoc + we want
predictable styling for headings, tables, and code spans.

Supports:
  - # / ## / ### / #### headings
  - paragraphs
  - bold (**...**), italic (*...*) and inline code (`...`)
  - bullet lists (-, *) and numbered lists (1. 2. ...)
  - markdown tables (| col | col |)
  - fenced code blocks (``` ... ```)
  - horizontal rules (---)
  - blockquotes (>)
"""
from __future__ import annotations
import re, sys, pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)")
TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{2,}.*\|")
LIST_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)")
LIST_NUMBER_RE = re.compile(r"^\s*\d+\.\s+(.*)")
HRULE_RE = re.compile(r"^---+\s*$")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)")


def add_run(para, text, *, bold=False, italic=False, code=False):
    if not text:
        return
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x8B, 0x4F, 0x00)


def render_inline(para, text):
    """Inline formatting: **bold**, *italic*, `code`."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            add_run(para, text[pos:m.start()])
        chunk = m.group(0)
        if chunk.startswith("**"):
            add_run(para, chunk[2:-2], bold=True)
        elif chunk.startswith("*"):
            add_run(para, chunk[1:-1], italic=True)
        elif chunk.startswith("`"):
            add_run(para, chunk[1:-1], code=True)
        pos = m.end()
    if pos < len(text):
        add_run(para, text[pos:])


def shade_cell(cell, color_hex):
    """Add a background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def parse_table(lines, start):
    """Parse a markdown table starting at lines[start]. Return (rows, end_idx)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.startswith("|"):
            break
        if TABLE_SEP_RE.match(line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def render(md_path: pathlib.Path, out_path: pathlib.Path):
    doc = Document()

    # Page margins — comfortable for printing
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code blocks
        if line.strip().startswith("```"):
            if in_code:
                # Close
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.4)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                # Light gray shading via pPr
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F4F4F4")
                p_pr.append(shd)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Horizontal rule
        if HRULE_RE.match(line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Heading
        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Strip trailing emoji/markdown
            heading_level = min(level, 4)
            h = doc.add_heading(level=heading_level)
            render_inline(h, text)
            for run in h.runs:
                if heading_level == 1:
                    run.font.size = Pt(20)
                    run.font.color.rgb = RGBColor(0x12, 0x36, 0x6E)
                elif heading_level == 2:
                    run.font.size = Pt(15)
                    run.font.color.rgb = RGBColor(0x1A, 0x4D, 0x90)
                elif heading_level == 3:
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x2C, 0x55, 0x9C)
                else:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # Blockquote
        m = BLOCKQUOTE_RE.match(line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(m.group(1))
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Table
        if line.startswith("|") and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1]):
            rows, end = parse_table(lines, i)
            if rows:
                cols = len(rows[0])
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row[:cols]):
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        if ri == 0:
                            shade_cell(cell, "1A4D90")
                            run = para.add_run(val)
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
                        else:
                            render_inline(para, val)
                            for r in para.runs:
                                r.font.size = Pt(10)
                doc.add_paragraph()
            i = end
            continue

        # Bullet list
        m = LIST_BULLET_RE.match(line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            render_inline(p, m.group(1))
            i += 1
            continue

        # Numbered list
        m = LIST_NUMBER_RE.match(line)
        if m:
            p = doc.add_paragraph(style="List Number")
            render_inline(p, m.group(1))
            i += 1
            continue

        # Empty line
        if line.strip() == "":
            i += 1
            continue

        # Default: paragraph
        p = doc.add_paragraph()
        render_inline(p, line)
        i += 1

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    src = pathlib.Path(sys.argv[1] if len(sys.argv) > 1 else "docs/TZ_FULLSTACK.md")
    dst = pathlib.Path(sys.argv[2] if len(sys.argv) > 2 else "docs/TZ_FULLSTACK.docx")
    out = render(src, dst)
    print(f"OK {out} ({out.stat().st_size:,} bytes)")
